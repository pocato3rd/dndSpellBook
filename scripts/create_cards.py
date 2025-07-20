import re
import glob
import os, pathlib
import math
from io import StringIO

from bs4 import BeautifulSoup
import pandas as pd
import numpy as np

# Imports to work with Word Documents
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Mm, Inches, RGBColor

# Imports to work with images
from PIL import Image, ImageDraw, ImageFont

import logging
from customLogFormatter import CustomFormatter

log = logging.getLogger("create_cards.py")
log.setLevel(logging.INFO)

cons_handler = logging.StreamHandler()
cons_handler.setLevel(logging.INFO)
cons_handler.setFormatter(CustomFormatter())
log.addHandler(cons_handler)


ROOT_DIR = pathlib.Path(__file__).parent.parent.resolve()

class Card:
    SCHOOL_COLORS = {
        "abjuration": "00b0f0",
        "conjuration": "ed7d31",
        "divination": "808080",
        "enchantment": "ff85ff",
        "evocation": "c00000",
        "illusion": "7030a0",
        "necromancy": "00b050",
        "transmutation": "833c0b"
    }

    TEMPLATE_COLOR = SCHOOL_COLORS['conjuration'].lower()

    CLASSES = ["Artificer", "Bard", "Cleric", "Druid", "Paladin", "Ranger", "Sorcerer", "Warlock", "Wizard"]
    REQUIREMENT_ORDER = ['concentration', 'ritual', 'verbal', 'somatic', 'material_comp']

    # Construction constants
    TABLE_ROW_LIMIT_PER_PAGE = 19
    SUPPORTED_FONT_SIZES = [8, 7, 6.5]
    MAX_FONT_WITH_2_PAGES = 0
    LINE_LIMITS = {
        # [1st page line limit, nth page line limit, chars/line]
        '8': [13, 26, 54],
        '7': [17, 28, 55],
        '6.5': [20, 32, 54],
    }

    def __init__(self, spell_row, output_dir="./outputs"):
        self.output_dir = output_dir

        # Initialized Required Spell Details
        self.name = spell_row.get("Spell Name").replace("/", "-")
        self.level = str(spell_row.get("Level"))

        self.school = spell_row.get("School")

        self.range = str(spell_row.get("Range"))
        self.duration = str(spell_row.get("Duration"))
        self.casting_time = str(spell_row.get("Casting Time"))

        # Boolean Toggles
        self.concentration = bool(spell_row['Concentration'])
        self.ritual = bool(spell_row["Ritual"])
        self.verbal = bool(spell_row["Verbal"])
        self.somatic = bool(spell_row["Somatic"])
        self.material = bool(spell_row["Material"])
        self.has_tables = bool(spell_row["Has Tables"])

        # Process supported classes
        self.supported_classes = dict()
        for c in Card.CLASSES:
            class_app = str(spell_row[c])
            if class_app.lower() in ['nan', 'no', '']:
                continue
            self.supported_classes[c] = class_app


        # Additional inputs
        self.description = spell_row.get("Description").split("|")
        self.source = spell_row.get("Source")

        # Optional fields
        self.set_material_components(spell_row.get("Material Component"))
        self.set_blurb(spell_row.get("Blurb"))

        # Start configuring the card based on fixed details.
        self.color = Card.SCHOOL_COLORS.get(self.school.lower(), "aaa").lower()
        self.description_len = 0

        for d in self.description: self.description_len += len(d)


    # Getters
    def get_has_tables(self):
        return self.has_tables

    def get_name(self):
        return self.name

    def get_level(self):
        return self.level
    
    def get_color(self):
        return self.color
    
    def get_classes(self):
        return self.supported_classes

    def get_class_info(self, class_name):
        return self.supported_classes.get(class_name, '').lower()
    
    def get_range(self):
        return self.range
    
    def get_duration(self):
        return self.duration
    
    def get_casting_time(self):
        return self.casting_time
    
    def get_description(self):
        return self.description

    def get_material_components(self):
        if self.material_comp is None:
            return ''
        else: 
            return self.material_comp
        
    def get_blurb(self):
        if self.short_blurb is None:
            return ''
        else:
            return self.short_blurb

    # setters
    def set_material_components(self, material_comp_text):
        if type(material_comp_text) in [np.float64, float] and np.isnan(material_comp_text):
            self.material_comp = None
        else:
            self.material_comp = material_comp_text

    def set_blurb(self, blurb_text):
        if type(blurb_text) in [np.float64, float] and np.isnan(blurb_text):
            self.short_blurb = None
        else:
            self.short_blurb = blurb_text

    # Saving Methods
    def save_as_img(self):
        """
        Save the card to an image file.
        """
        pass

    def save_as_docx(self):
        """
        Save the card to a docx file.
        """

        use_font_size, expected_page_count = self.get_font_size_and_page_count()

         # Get the template docx
        document = Document(os.path.join(ROOT_DIR,'resources','template_cards','TEMPLATE.docx'))
        
        # Add the style types
        styles = document.styles

        style = styles.add_style('Description', WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(use_font_size)

        style = styles.add_style('Line Break', WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(4)

        style = styles.add_style('Table Description', WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles['Description']
        style.font.size = Pt(5.5)


        # Update the spell requirement images
        for i in range(len(Card.REQUIREMENT_ORDER)):
            inline_elem = document.inline_shapes[i]
            req_type = Card.REQUIREMENT_ORDER[i]

            if req_type == "material_comp":
                material_comp_bool = self.material_comp is not None
                with open(os.path.join(ROOT_DIR,f'./resources/images/{req_type}/{str(material_comp_bool).lower()}.png'), 'br') as f:
                    img_bytes = f.read()
            elif req_type == "concentration":
                with open(os.path.join(ROOT_DIR,f'./resources/images/{req_type}/{str(self.concentration).lower()}.png'), 'br') as f:
                    img_bytes = f.read()
            elif req_type == "ritual":
                with open(os.path.join(ROOT_DIR,f'./resources/images/{req_type}/{str(self.ritual).lower()}.png'), 'br') as f:
                    img_bytes = f.read()
            elif req_type == "verbal":
                with open(os.path.join(ROOT_DIR,f'./resources/images/{req_type}/{str(self.verbal).lower()}.png'), 'br') as f:
                    img_bytes = f.read()
            elif req_type == "somatic":
                with open(os.path.join(ROOT_DIR,f'./resources/images/{req_type}/{str(self.somatic).lower()}.png'), 'br') as f:
                    img_bytes = f.read()

            rId = inline_elem._inline.graphic.graphicData.pic.blipFill.blip.embed
            document.part.related_parts[rId]._blob = img_bytes

        # Spell Name, the 0th row and cell of each table
        for i, t in enumerate(document.tables):
            spell_name_elem = t.rows[0].cells[0]

            updated_text = False
            for t_elem in spell_name_elem._tc.iterdescendants(qn('w:t')):
                if not updated_text:
                    if expected_page_count > 1 or self.get_has_tables(): 
                        # there are at least two pages
                        t_elem.text = f'{self.get_name()} (Part {i+1})'
                    else:
                        t_elem.text = f'{self.get_name()}'
                    
                    updated_text = True
                else: t_elem.text = ''
        # Spell Level (0th table, 0th row, 2nd cell, w:t element)
        spell_level_elem = document.tables[0].rows[0].cells[2]
        spell_level_text = next(spell_level_elem._tc.iterdescendants(qn('w:t')))
        spell_level_text.text = self.get_level()
        spell_level_bottom = next(spell_level_elem._tc.iterdescendants(qn('w:bottom')))
        spell_level_bottom.attrib[qn('w:color')] = self.get_color()

        # Range
        range_elem = document.tables[0].rows[1].cells[1]
        for i, r_elem in enumerate(range_elem._tc.iterdescendants(qn('w:r'))):
            if i == 0:
                r_elem.first_child_found_in('w:t').text = self.get_range()
            else:
                r_elem.first_child_found_in('w:t').text = ''

        # Duration
        duration_elem = document.tables[0].rows[2].cells[1]
        for i, r_elem in enumerate(duration_elem._tc.iterdescendants(qn('w:r'))):
            if i == 0:
                r_elem.first_child_found_in('w:t').text = self.get_duration()
            else:
                r_elem.first_child_found_in('w:t').text = ''

        # Casting Time
        casting_time_elem = document.tables[0].rows[3].cells[1]
        for i, r_elem in enumerate(casting_time_elem._tc.iterdescendants(qn('w:r'))):
            if i == 0:
                r_elem.first_child_found_in('w:t').text = self.get_casting_time()
            else:
                r_elem.first_child_found_in('w:t').text = ''

        # Material Components (if applicable)
        material_comp_elem = document.tables[0].rows[5].cells[1]
        document.tables[0].rows[5].height = Mm(4.8)
        paragraph_elem = material_comp_elem._tc.first_child_found_in('w:p')
        for i, r_elem in enumerate(material_comp_elem._tc.iterdescendants(qn('w:r'))):
            if i == 0:
                r_elem.first_child_found_in('w:t').text = self.get_material_components()
            else:
                r_elem.first_child_found_in('w:t').text = ''

        # Short Blurb (if applicable)
        short_blurb_elem = document.tables[0].rows[6].cells[1]
        paragraph_elem = short_blurb_elem._tc.first_child_found_in('w:p')
        for i, r_elem in enumerate(short_blurb_elem._tc.iterdescendants(qn('w:r'))):
            if i == 0:
                r_elem.first_child_found_in('w:t').text = self.get_blurb()
            else:
                r_elem.first_child_found_in('w:t').text = ''

        # Blanket update to the background colors
        for shd_prop in document.element.iterdescendants(qn('w:shd')):
            if shd_prop.attrib.get(qn('w:fill'), '').lower() == Card.TEMPLATE_COLOR:
                shd_prop.attrib[qn('w:fill')] = self.get_color()

                if shd_prop.attrib.get(qn('w:themeFill')):
                    shd_prop.attrib.pop(qn('w:themeFill'))

        # Blanket update to the table borders
        for tbProps in document.element.iterdescendants(qn('w:tcBorders')):
                for borderItem in tbProps.iterchildren():
                    if borderItem.attrib.get(qn('w:color'), '').lower() in [Card.TEMPLATE_COLOR, 'ff85ff']:
                        borderItem.attrib[qn('w:color')] = self.get_color()

                        if borderItem.attrib.get(qn('w:themeColor')):
                            borderItem.attrib.pop(qn('w:themeColor'))

        # Update the class color-coding
        for t_Tb in document.tables:
            for row in t_Tb.rows:
                for c in row.cells:
                    # Found the class list, update it based on the spell
                    # class list is in table 0, rows [1,2,3,4], cell 2
                    if "artificer" in c.text.lower(): 
                        for cell_child in c._tc.getchildren():
                            current_dnd_class = None

                            r_elem = cell_child.first_child_found_in('w:r')
                            if r_elem is not None:
                                text_elem = r_elem.first_child_found_in('w:t')
                                
                                # determine which class are we working with
                                if text_elem is not None:
                                    current_dnd_class = text_elem.text.strip()

                                # now that we have a class, we need to check if it's applicable (color it) and optional (underline)
                                rPr_elem = r_elem.first_child_found_in('w:rPr')

                                # remove the underline if it exists for a fresh start
                                underline_elem = rPr_elem.first_child_found_in('w:u')
                                if underline_elem is not None:
                                    rPr_elem.remove(underline_elem)

                                if current_dnd_class in self.get_classes():
                                    current_color = self.get_color()
                                    if self.get_class_info(current_dnd_class) == 'optional':
                                        # add in the underline
                                        u_elem = OxmlElement('w:u')
                                        u_elem.set(qn('w:val'), 'single')
                                        rPr_elem.append(u_elem)
                                else:
                                    # reset the color
                                    current_color = "000000"

                                color_elem = r_elem.first_child_found_in('w:rPr').first_child_found_in('w:color')
                                color_elem.attrib[qn('w:val')] = current_color
                                # remove the theme color if it exists so we can modify the color directly
                                if color_elem.attrib.get(qn('w:themeColor')):
                                    color_elem.attrib.pop(qn('w:themeColor'))

        # Update descriptions
        # last row of each table, cell 0
        description_cell = document.tables[0].rows[-1].cells[0]
        document.tables[0].rows[-1].height = Inches(1.85)
        document.tables[1].rows[-1].height = Inches(3.05)
        old_paragraphs = description_cell.paragraphs

        current_line = 0
        line_limit = Card.LINE_LIMITS[str(use_font_size)][0]
        page_count = 0

        for i, d in enumerate(self.get_description()):

            if current_line + math.ceil(len(d)/Card.LINE_LIMITS[str(use_font_size)][2]) > line_limit:
                # this would exceed the page, put it on the next one
                for p in old_paragraphs: description_cell._tc.remove(p._element)
                
                # remove the unnecessary space
                if description_cell.paragraphs:
                    description_cell._tc.remove(description_cell.paragraphs[-1]._element)

                page_count += 1

                if page_count == len(document.tables):
                    template_tbl = document.tables[1]._tbl.__deepcopy__(None)

                    template_break = document.tables[-1]._element.getnext()
                    template_break.addnext(template_tbl)
                    template_tbl.addnext(template_break.__deepcopy__(None))

                    # document.tables[-1]._element.getnext().addnext(template_tbl)
                    # template_tbl._element.addnext()

                    # new_table = document.tables[-1]._tbl.addnext(template_tbl)

                    document.tables.append(template_tbl)

                    spell_name_elem = next(document.tables[-1].rows[0].cells[0]._tc.iterdescendants(qn('w:t')))
                    spell_name_elem.text = f'{self.get_name()} (Part {page_count+1})'

                description_cell = document.tables[page_count].rows[-1].cells[0]
                document.tables[1].rows[-1].height = Inches(3.05)
                old_paragraphs = description_cell.paragraphs

                # go to next page, increase current_limit
                current_line = 0
                line_limit = Card.LINE_LIMITS[str(use_font_size)][1]     

            # Time to actually add the description paragraph
            # First remove <p> and </p> tags:
            use_d = d[:]
            while "<p>" in use_d: use_d = use_d.replace('<p>', '')
            while "</p>" in use_d: use_d = use_d.replace('</p>', '')

            if '<li>' in use_d:            
                # unordered list
                use_d = re.sub('<ul>?<li>', '\u2022 ', use_d)
                use_d = re.sub('</li>?</ul>', '', use_d)

                # ordered list
                use_d = re.sub('<ol>?<li>', '\u2022 ', use_d)
                use_d = re.sub('</li>?</ol>', '', use_d)

            paragraph_start_idx = 0
            runs_to_add = [] # list of tuples (IS_BOLD, str)

            ## First check for bold or strong tags
            if "<strong>" in use_d or "<b>" in use_d:
                # find text data between <strong> and </strong>
                x = re.finditer("(<strong>.*</strong>)|(<b>.*</b>)", use_d)

                for match in x:
                    match_start, match_end = match.span()

                    if match_start != paragraph_start_idx:
                        # we have non-bolded text
                        runs_to_add.append((False, use_d[paragraph_start_idx:match_start]))
                    if "<b>" in match.group():
                        # <b> was used to bold
                        runs_to_add.append((True, use_d[paragraph_start_idx+3:match_end-4]))
                        paragraph_start_idx=match_end+1
                    elif "<strong>" in match.group():
                        # <strong> was used to bold
                        runs_to_add.append((True, use_d[paragraph_start_idx+8:match_end-9]))
                        paragraph_start_idx=match_end+1
            runs_to_add.append((False, use_d[paragraph_start_idx:]))
                    
            # Add the runs in
            p = description_cell.add_paragraph(style=styles['Description'])
            for is_bold, r in runs_to_add:
                # remove any extra tags (end tags, then start tags)
                use_r = r[:]
                for end_tag_match in re.finditer('</.*>', use_r):
                    use_r = re.sub(end_tag_match.group(),'', use_r)
                for start_tag_match in re.finditer('<.*>', use_r):
                    use_r = re.sub(start_tag_match.group(),'', use_r)
                
                runner = p.add_run(use_r+' ')
                if is_bold: runner.bold = True

            # if i < len(spell_details['description'])-1:
            description_cell.add_paragraph(' ', styles['Line Break'])

            current_line += math.ceil(len(d)/Card.LINE_LIMITS[str(use_font_size)][2]) + 1

        for p in old_paragraphs:
            description_cell._tc.remove(p._element)

        if page_count == 0:
            # did not use the second page of the template for descriptions
            cell = document.tables[-1].rows[-1].cells[0]
            for p in cell.paragraphs:
                p._parent._tc.remove(p._element)
        
        if self.get_has_tables():
            # Add in description table if applicable
            total_rows = 0
            for i, table_html in enumerate(glob.glob(os.path.join(ROOT_DIR,f'./resources/tables/{self.get_name()}_table*.html'))):
                tables = parse_html_table_into_py(table_html)
                row_count = tables[0].shape[0]

                # Create a new page if you haven't already
                if page_count == 0:
                    # No need to make a new page, this one is blank
                    cell = document.tables[-1].rows[-1].cells[0]
                    page_count += 1

                    spell_name_elem = next(document.tables[-1].rows[0].cells[0]._tc.iterdescendants(qn('w:t')))
                    spell_name_elem.text = f'{self.get_name()} (Part {page_count+1})'
                    document.tables[1].rows[-1].height = Inches(3.05)

                elif (total_rows + row_count <= Card.TABLE_ROW_LIMIT_PER_PAGE) and i > 0:
                    # no need to make an additional page
                    cell = document.tables[-1].rows[-1].cells[0]
                    pass
                
                else:
                    # Need to make a new page...
                    total_rows = 0
                    page_count += 1
                    template_tbl = document.tables[-1]._tbl.__deepcopy__(None)
                    template_break = document.tables[-1]._element.getnext()
                    template_break.addnext(template_tbl)
                    template_tbl.addnext(template_break.__deepcopy__(None))

                    document.tables.append(template_tbl)

                    spell_name_elem = next(document.tables[-1].rows[0].cells[0]._tc.iterdescendants(qn('w:t')))
                    spell_name_elem.text = f'{self.get_name()} (Part {page_count+1})'

                    cell = document.tables[page_count].rows[-1].cells[0]
                    document.tables[-1].rows[-1].height = Inches(3.05)

                    for p in cell.paragraphs: 
                        cell._tc.remove(p._element)
                    for t in cell.tables:
                        cell._tc.remove(t._element)

                cell.add_paragraph('',styles['Line Break'])

                new_table = add_table_into_docx(tables, parent=document.tables[-1].rows[-1].cells[0], styles=styles, school_color=self.get_color())
                total_rows += row_count

        if page_count == 0:
            # we didn't need a second page, remove it
            document.tables[1]._element.getparent().remove(document.tables[1]._element)

        # with open(os.path.join(ROOT_DIR,"logs/document.xml"), "w", encoding='utf-8') as f:
        #     f.write(document.element.xml)
        document.save(self.get_output_location(docx=True))



    # Construction helpers
    def number_of_pages(self, font_size):
        # determine number of lines used by descriptions based on the font size
        current_line = 0
        page_num = 0
        line_limit = Card.LINE_LIMITS[str(font_size)][0]
        chars_per_line = Card.LINE_LIMITS[str(font_size)][2]

        for d in self.description:
            # if we'd exceed the current limit, we need a new page
            if current_line + math.ceil(len(d)/chars_per_line) > line_limit:
                current_line = 0
                page_num += 1
                line_limit = Card.LINE_LIMITS[str(font_size)][1]

            current_line += math.ceil(len(d)/chars_per_line) + 1

        # +1 due to zero index
        return page_num + 1        
    
    def get_font_size_and_page_count(self):

        max_font_with_2_pages = Card.MAX_FONT_WITH_2_PAGES

        for use_font_size in Card.SUPPORTED_FONT_SIZES:
            expected_page_count = self.number_of_pages(use_font_size)

            if expected_page_count == 1:
                # this font size works
                break
            elif expected_page_count == 2:
                max_font_with_2_pages = max(max_font_with_2_pages, use_font_size)
                continue
        
        if expected_page_count > 1 and max_font_with_2_pages != 0:
            use_font_size = max_font_with_2_pages
        
        return use_font_size, expected_page_count

    def get_output_location(self, docx=True) -> str:
        to_return = f'{self.output_dir}/level_{self.get_level()}'
        os.makedirs(to_return, exist_ok=True)

        to_return += f'/{self.get_name()}' 
        if docx:
            to_return += '.docx'
        else:
            # assume .png
            to_return += '.png'

        return to_return


def number_of_pages(descriptions, font_size):
    # determine number of lines used by descriptions based on font size
    current_line = 0
    page_num = 0
    line_limit = Card.LINE_LIMITS[str(font_size)][0]
    chars_per_line = Card.LINE_LIMITS[str(font_size)][2]

    for d in descriptions:
        # if we'd exceed the current limit, we need a new page
        if current_line + math.ceil(len(d)/chars_per_line) > line_limit:
            current_line = 0
            page_num += 1
            line_limit = Card.LINE_LIMITS[str(font_size)][1]

        current_line += math.ceil(len(d)/chars_per_line) + 1

    # +1 due to zero index
    return page_num + 1


def parse_html_table_into_py(table_html):
    # Parse the saved html table into usable Python data structures
    # Returns 
    #   * table_headers, (2d) ndarray of bools, access by [row, col]
    #   * table_contents, 2d nested lists of strings, access by [row][col]
    #   * table_row_span, (2d) ndarray of ints, access by [row, col]
    #   * table_col_span, (2d) ndarray of ints, access by [row, col]

    # read the html
    with open(table_html, 'r', encoding='utf-8') as f:
        soup = BeautifulSoup(f.read(), "html.parser")

    # get the table into usable high level structures
    table = soup.find('table')
    df = pd.read_html(StringIO(table.prettify()), header=None)[0]

    # dimensions
    row_count = len(table.find_all('tr'))
    col_count = df.shape[1]
    shape = (row_count, col_count)

    # make our data structures
    table_headers = np.ndarray(shape, dtype=bool)
    table_row_span = np.zeros(shape,dtype=int)
    table_col_span = np.zeros(shape,dtype=int)
    # ndarray is weird with strings
    table_contents = [None]*row_count

    for i, soup_row in enumerate(table.find_all('tr')):
        if soup_row.get('rowspan',None):
            log.warning(f"Detected unhandled rowspan in {table_html}")
        
        row_span = 0
        use_j = 0
        table_contents[i] = [None]*col_count

        for soup_col in soup_row.find_all(['th', 'td']):
            is_header = soup_col.name == 'th'

            cell_contents = [] # (content, bold, italics)

            # get each tag in table cell (including untagged spans)
            for c in soup_col.contents:
                # remove newlines
                bold = 'strong' in str(c)
                italic = 'em' in str(c)
                
                use_c = str(c).replace('\n','')
                # remove bolds
                use_c = use_c.replace('<strong>','').replace('</strong>','')
                # remove italics
                use_c = use_c.replace('<em>','').replace('</em>','')

                if '<' in use_c:
                    # there were unhandled tags, remove them
                    use_c = re.sub('</.*>','', use_c)
                    use_c = re.sub('<.*>','', use_c)

                # strip
                use_c = use_c.strip()

                if use_c:
                    cell_contents.append((use_c, bold, italic))

            col_span = soup_col.get('colspan', 1)
            
            table_headers[i, use_j] = is_header
            table_contents[i][use_j] = cell_contents[:]
            table_row_span[i,use_j] = int(row_span)
            table_col_span[i,use_j] = int(col_span)
            
            use_j += int(col_span)

    return table_headers, table_contents, table_row_span, table_col_span

def add_table_into_docx(py_tables, parent, styles, school_color):    
    
    # Now that we have the headers, span locations, merging style, and innerHtmls
    # Create the docx table cell by cell
    table_headers, table_contents, table_row_span, table_col_span = py_tables

    # get the dimensions and make the table
    row_count, col_count = table_headers.shape
    docx_table = parent.add_table(rows=row_count, cols=col_count)

    # an empty paragraph is created for some reason, remove it
    next_p = docx_table._element.getnext()
    if next_p.tag[-1] == "p":
        parent._tc.remove(next_p)

    # set the table style to something basic
    docx_table.style = 'Table Grid'
    docx_table.autofit = False
    docx_table.allow_autofit = False
    docx_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # width to span most of the card
    table_width = Inches(2.3).emu

    # remove the autofit type and set the tables width directly
    tblW = docx_table._tblPr.getchildren()[1]
    tblW.attrib.pop(qn("w:type"))
    tblW.attrib[qn("w:w")] = str(table_width)

    # remove the grid with gridCol because it was not adjusting size
    grid = list(docx_table._element.iterchildren(qn('w:tblGrid')))[0]
    grid.getparent().remove(grid)

    # populate all data
    for i in range(row_count):
        row = docx_table.rows[i]
        row.height=Inches(0.1)

        for j in range(col_count):
            cell_content = table_contents[i][j]
            cell = docx_table.rows[i].cells[j]
            cell._tc.autofit=False
            cell._tc.width=str(int(table_width/col_count))

            # set the text style
            current_paragraph = cell.paragraphs[0]
            current_paragraph.style = 'Table Description'

            if table_headers[i,j]:
                # this is a header, color it as such
                tcPr = cell._tc.tcPr
                shd_elem = OxmlElement('w:shd')
                shd_elem.set(qn('w:fill'), school_color)
                shd_elem.set(qn('w:val'), 'clear')
                tcPr.append(shd_elem)

            if cell_content is not None:
                # for each un/tagged span, make a run with the right properties
                for c in cell_content:
                    runner = current_paragraph.add_run(c[0]+" ")
                    if c[1]: runner.bold = True
                    if c[2]: runner.italic = True
                    
                    if table_headers[i,j]:
                        # assume this is okay contrast
                        runner.font.color.rgb = RGBColor(0xff,0xff,0xff)

    # go through and merge cells across colspans
    for i in range(row_count):
        for j in range(col_count):
            colspan = table_col_span[i,j]
            if colspan > 1:
                # need to merge cells
                for span_idx in range(colspan-1):
                    docx_table.rows[i].cells[j].merge(docx_table.rows[i].cells[j+span_idx+1])

    return docx_table


def create_spell_card(spell_details, output_loc):

    description_char_count = 0
    for d in spell_details['description']: description_char_count += len(d)
    spell_details['description_length'] = description_char_count

    supported_font_sizes = [8, 7, 6.5]
    max_font_with_2 = 0

    for use_font_size in supported_font_sizes:
        expected_page_count = number_of_pages(spell_details['description'], use_font_size)
        if expected_page_count == 1: break
        elif expected_page_count == 2:
            max_font_with_2 = max(max_font_with_2, use_font_size)
            continue
    if expected_page_count > 1 and max_font_with_2 != 0: 
        use_font_size = max_font_with_2 
        
    # The color to set based on the spell
    use_color = Card.SCHOOL_COLORS[spell_details['school']].lower()

    # The color used in the template
    look_for_color = Card.SCHOOL_COLORS['conjuration'].lower()

    # Get the template docx
    document = Document(os.path.join(ROOT_DIR,'resources','template_cards','TEMPLATE.docx'))
    
    # Add the style types
    styles = document.styles

    style = styles.add_style('Description', WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(use_font_size)

    style = styles.add_style('Line Break', WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(4)

    style = styles.add_style('Table Description', WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles['Description']
    style.font.size = Pt(5.5)

    # save the document xml before making changes 
    # with open(os.path.join(ROOT_DIR,"./logs/pre-document.xml"), "w", encoding='utf-8') as f:
    #     f.write(document.element.xml)

    # Update the spell requirement images
    for i in [0,1,2,3,4]:
        inline_elem = document.inline_shapes[i]
        req_type = Card.REQUIREMENT_ORDER[i]

        if req_type == "material_comp":
            material_comp_bool = ("material_comp" in spell_details)
            with open(os.path.join(ROOT_DIR,f'./resources/images/{req_type}/{str(material_comp_bool).lower()}.png'), 'br') as f:
                img_bytes = f.read()

        else:
            with open(os.path.join(ROOT_DIR,f'./resources/images/{req_type}/{str(spell_details[req_type]).lower()}.png'), 'br') as f:
                img_bytes = f.read()

        rId = inline_elem._inline.graphic.graphicData.pic.blipFill.blip.embed
        document.part.related_parts[rId]._blob = img_bytes

    # Spell Name, the 0th row and cell of each table
    for i, t in enumerate(document.tables):
        spell_name_elem = t.rows[0].cells[0]

        updated_text = False
        for t_elem in spell_name_elem._tc.iterdescendants(qn('w:t')):
            if not updated_text:
                if expected_page_count > 1 or spell_details.get('has_tables', False): 
                    # there are at least two pages
                    t_elem.text = f'{spell_details["name"]} (Part {i+1})'
                else:
                    t_elem.text = f'{spell_details["name"]}'
                
                updated_text = True
            else: t_elem.text = ''
        
    # Spell Level (0th table, 0th row, 2nd cell, w:t element)
    spell_level_elem = document.tables[0].rows[0].cells[2]
    spell_level_text = next(spell_level_elem._tc.iterdescendants(qn('w:t')))
    spell_level_text.text = spell_details["level"]
    spell_level_bottom = next(spell_level_elem._tc.iterdescendants(qn('w:bottom')))
    spell_level_bottom.attrib[qn('w:color')] = use_color

    # Range
    range_elem = document.tables[0].rows[1].cells[1]
    for i, r_elem in enumerate(range_elem._tc.iterdescendants(qn('w:r'))):
        if i == 0:
            r_elem.first_child_found_in('w:t').text = spell_details['range']
        else:
            r_elem.first_child_found_in('w:t').text = ''

    # Duration
    duration_elem = document.tables[0].rows[2].cells[1]
    for i, r_elem in enumerate(duration_elem._tc.iterdescendants(qn('w:r'))):
        if i == 0:
            r_elem.first_child_found_in('w:t').text = spell_details['duration']
        else:
            r_elem.first_child_found_in('w:t').text = ''

    # Casting Time
    casting_time_elem = document.tables[0].rows[3].cells[1]
    for i, r_elem in enumerate(casting_time_elem._tc.iterdescendants(qn('w:r'))):
        if i == 0:
            r_elem.first_child_found_in('w:t').text = spell_details['casting_time']
        else:
            r_elem.first_child_found_in('w:t').text = ''

    # Material Components (if applicable)
    material_comp_elem = document.tables[0].rows[5].cells[1]
    document.tables[0].rows[5].height = Mm(4.8)
    paragraph_elem = material_comp_elem._tc.first_child_found_in('w:p')
    for i, r_elem in enumerate(material_comp_elem._tc.iterdescendants(qn('w:r'))):
        if i == 0:
            r_elem.first_child_found_in('w:t').text = spell_details.get('material_comp', ' ')
        else:
            r_elem.first_child_found_in('w:t').text = ''

    # Short Blurb (if applicable)
    short_blurb_elem = document.tables[0].rows[6].cells[1]
    paragraph_elem = short_blurb_elem._tc.first_child_found_in('w:p')
    for i, r_elem in enumerate(short_blurb_elem._tc.iterdescendants(qn('w:r'))):
        if i == 0:
            r_elem.first_child_found_in('w:t').text = spell_details.get('short_blurb', '')
        else:
            r_elem.first_child_found_in('w:t').text = ''

    # Blanket update to the background colors
    for shd_prop in document.element.iterdescendants(qn('w:shd')):
        if shd_prop.attrib.get(qn('w:fill'), '').lower() == look_for_color:
            shd_prop.attrib[qn('w:fill')] = use_color

            if shd_prop.attrib.get(qn('w:themeFill')):
                shd_prop.attrib.pop(qn('w:themeFill'))

    # Blanket update to the table borders
    for tbProps in document.element.iterdescendants(qn('w:tcBorders')):
            for borderItem in tbProps.iterchildren():
                if borderItem.attrib.get(qn('w:color'), '').lower() in [look_for_color, 'ff85ff']:
                    borderItem.attrib[qn('w:color')] = use_color

                    if borderItem.attrib.get(qn('w:themeColor')):
                        borderItem.attrib.pop(qn('w:themeColor'))


    # Update the class color-coding
    for t_Tb in document.tables:
        for row in t_Tb.rows:
            for c in row.cells:
                # Found the class list, update it based on the spell
                # class list is in table 0, rows [1,2,3,4], cell 2
                if "artificer" in c.text.lower(): 
                    for cell_child in c._tc.getchildren():
                        current_dnd_class = None

                        r_elem = cell_child.first_child_found_in('w:r')
                        if r_elem is not None:
                            text_elem = r_elem.first_child_found_in('w:t')
                            
                            # determine which class are we working with
                            if text_elem is not None:
                                current_dnd_class = text_elem.text.strip()

                            # now that we have a class, we need to check if it's applicable (color it) and optional (underline)
                            rPr_elem = r_elem.first_child_found_in('w:rPr')

                            # remove the underline if it exists for a fresh start
                            underline_elem = rPr_elem.first_child_found_in('w:u')
                            if underline_elem is not None:
                                rPr_elem.remove(underline_elem)

                            if current_dnd_class in spell_details['applicable_classes']:
                                current_color = use_color
                                if spell_details['applicable_classes'][current_dnd_class].lower() == 'optional':
                                    # add in the underline
                                    u_elem = OxmlElement('w:u')
                                    u_elem.set(qn('w:val'), 'single')
                                    rPr_elem.append(u_elem)
                            else:
                                # reset the color
                                current_color = "000000"

                            color_elem = r_elem.first_child_found_in('w:rPr').first_child_found_in('w:color')
                            color_elem.attrib[qn('w:val')] = current_color
                            # remove the theme color if it exists so we can modify the color directly
                            if color_elem.attrib.get(qn('w:themeColor')):
                                color_elem.attrib.pop(qn('w:themeColor'))

    # Update descriptions
    # last row of each table, cell 0
    description_cell = document.tables[0].rows[-1].cells[0]
    document.tables[0].rows[-1].height = Inches(1.85)
    document.tables[1].rows[-1].height = Inches(3.05)
    old_paragraphs = description_cell.paragraphs

    current_line = 0
    line_limit = Card.LINE_LIMITS[str(use_font_size)][0]
    page_count = 0

    for i, d in enumerate(spell_details['description']):

        if current_line + math.ceil(len(d)/Card.LINE_LIMITS[str(use_font_size)][2]) > line_limit:
            # this would exceed the page, put it on the next one
            for p in old_paragraphs: description_cell._tc.remove(p._element)
            
            # remove the unnecessary space
            if description_cell.paragraphs:
                description_cell._tc.remove(description_cell.paragraphs[-1]._element)

            page_count += 1

            if page_count == len(document.tables):
                template_tbl = document.tables[1]._tbl.__deepcopy__(None)

                template_break = document.tables[-1]._element.getnext()
                template_break.addnext(template_tbl)
                template_tbl.addnext(template_break.__deepcopy__(None))

                # document.tables[-1]._element.getnext().addnext(template_tbl)
                # template_tbl._element.addnext()

                # new_table = document.tables[-1]._tbl.addnext(template_tbl)

                document.tables.append(template_tbl)

                spell_name_elem = next(document.tables[-1].rows[0].cells[0]._tc.iterdescendants(qn('w:t')))
                spell_name_elem.text = f'{spell_details["name"]} (Part {page_count+1})'

            description_cell = document.tables[page_count].rows[-1].cells[0]
            document.tables[1].rows[-1].height = Inches(3.05)
            old_paragraphs = description_cell.paragraphs

            # go to next page, increase current_limit
            current_line = 0
            line_limit = Card.LINE_LIMITS[str(use_font_size)][1]     

        # Time to actually add the description paragraph
        # First remove <p> and </p> tags:
        use_d = d[:]
        while "<p>" in use_d: use_d = use_d.replace('<p>', '')
        while "</p>" in use_d: use_d = use_d.replace('</p>', '')

        if '<li>' in use_d:            
            # unordered list
            use_d = re.sub('<ul>?<li>', '\u2022 ', use_d)
            use_d = re.sub('</li>?</ul>', '', use_d)

            # ordered list
            use_d = re.sub('<ol>?<li>', '\u2022 ', use_d)
            use_d = re.sub('</li>?</ol>', '', use_d)

        paragraph_start_idx = 0
        runs_to_add = [] # list of tuples (IS_BOLD, str)

        ## First check for bold or strong tags
        if "<strong>" in use_d or "<b>" in use_d:
            # find text data between <strong> and </strong>
            x = re.finditer("(<strong>.*</strong>)|(<b>.*</b>)", use_d)

            for match in x:
                match_start, match_end = match.span()

                if match_start != paragraph_start_idx:
                    # we have non-bolded text
                    runs_to_add.append((False, use_d[paragraph_start_idx:match_start]))
                if "<b>" in match.group():
                    # <b> was used to bold
                    runs_to_add.append((True, use_d[paragraph_start_idx+3:match_end-4]))
                    paragraph_start_idx=match_end+1
                elif "<strong>" in match.group():
                    # <strong> was used to bold
                    runs_to_add.append((True, use_d[paragraph_start_idx+8:match_end-9]))
                    paragraph_start_idx=match_end+1
        runs_to_add.append((False, use_d[paragraph_start_idx:]))
                
        # Add the runs in
        p = description_cell.add_paragraph(style=styles['Description'])
        for is_bold, r in runs_to_add:
            # remove any extra tags (end tags, then start tags)
            use_r = r[:]
            for end_tag_match in re.finditer('</.*>', use_r):
                use_r = re.sub(end_tag_match.group(),'', use_r)
            for start_tag_match in re.finditer('<.*>', use_r):
                use_r = re.sub(start_tag_match.group(),'', use_r)
            
            runner = p.add_run(use_r+' ')
            if is_bold: runner.bold = True

        # if i < len(spell_details['description'])-1:
        description_cell.add_paragraph(' ', styles['Line Break'])

        current_line += math.ceil(len(d)/Card.LINE_LIMITS[str(use_font_size)][2]) + 1

    for p in old_paragraphs:
        description_cell._tc.remove(p._element)

    if page_count == 0:
        # did not use the second page of the template for descriptions
        cell = document.tables[-1].rows[-1].cells[0]
        for p in cell.paragraphs:
            p._parent._tc.remove(p._element)
    
    if spell_details.get('has_tables',False):
        # Add in description table if applicable
        total_rows = 0
        for i, table_html in enumerate(glob.glob(os.path.join(ROOT_DIR,f'./resources/tables/{spell_details["name"]}_table*.html'))):
            tables = parse_html_table_into_py(table_html)
            row_count = tables[0].shape[0]

            # Create a new page if you haven't already
            if page_count == 0:
                # No need to make a new page, this one is blank
                cell = document.tables[-1].rows[-1].cells[0]
                page_count += 1

                spell_name_elem = next(document.tables[-1].rows[0].cells[0]._tc.iterdescendants(qn('w:t')))
                spell_name_elem.text = f'{spell_details["name"]} (Part {page_count+1})'
                document.tables[1].rows[-1].height = Inches(3.05)

            elif (total_rows + row_count <= Card.TABLE_ROW_LIMIT_PER_PAGE) and i > 0:
                # no need to make an additional page
                cell = document.tables[-1].rows[-1].cells[0]
                pass
            
            else:
                # Need to make a new page...
                total_rows = 0
                page_count += 1
                template_tbl = document.tables[-1]._tbl.__deepcopy__(None)
                template_break = document.tables[-1]._element.getnext()
                template_break.addnext(template_tbl)
                template_tbl.addnext(template_break.__deepcopy__(None))

                document.tables.append(template_tbl)

                spell_name_elem = next(document.tables[-1].rows[0].cells[0]._tc.iterdescendants(qn('w:t')))
                spell_name_elem.text = f'{spell_details["name"]} (Part {page_count+1})'

                cell = document.tables[page_count].rows[-1].cells[0]
                document.tables[-1].rows[-1].height = Inches(3.05)

                for p in cell.paragraphs: 
                    cell._tc.remove(p._element)
                for t in cell.tables:
                    cell._tc.remove(t._element)

            cell.add_paragraph('',styles['Line Break'])

            new_table = add_table_into_docx(tables, parent=document.tables[-1].rows[-1].cells[0], styles=styles, school_color=use_color)
            total_rows += row_count

    if page_count == 0:
        # we didn't need a second page, remove it
        document.tables[1]._element.getparent().remove(document.tables[1]._element)

    # with open(os.path.join(ROOT_DIR,"logs/document.xml"), "w", encoding='utf-8') as f:
    #     f.write(document.element.xml)
    document.save(output_loc)

def parse_input_xlsx(input_xlsx):
    df = pd.read_excel(input_xlsx, sheet_name='Sheet1')
    filtered_df = df[df['Generate Card']]
    return filtered_df

def create_filtered_cards(df, output_dir):
    count_created = 1
    total_count = df.shape[0]
    spells_with_tables = set()

    for _, row in df.iterrows():

        if True:
            # use Card class
            newCard = Card(row, output_dir)
            log.info(f"[{count_created}/{total_count}]: Level {newCard.get_level()} spell, '{newCard.get_name()}' - generating...")
            newCard.save_as_docx()
            log.debug('done.')
            
            if newCard.get_has_tables():
                spells_with_tables.add(f"(Lvl {newCard.get_level()}) {newCard.get_name()}")
            count_created += 1

        else:
            # use standalone functions
            spell_details = {
                "name": row['Spell Name'],
                "level": str(row['Level']),
                "school": row['School'].lower(),
                "applicable_classes": dict(),
                "range": str(row['Range']),
                "duration": str(row['Duration']),
                "casting_time": str(row['Casting Time']),
                "material_comp": str(row['Material Component']),
                "concentration": bool(row['Concentration']),
                "ritual": bool(row["Ritual"]),
                "verbal": bool(row["Verbal"]),
                "somatic": bool(row["Somatic"]),
                "material": bool(row["Material"]),
                "description": row["Description"].split('|'),
                "has_tables": bool(row["Has Tables"]),
                "source": row["Source"],
                "short_blurb": str(row["Blurb"])
            }

            for c in Card.CLASSES:
                class_applicability = str(row[c])
                if class_applicability.lower() in ['nan', 'no']:
                    continue
                spell_details['applicable_classes'][c] = str(row[c])

            if spell_details['material_comp'].lower() == 'nan':
                spell_details.pop('material_comp')        
            if spell_details['short_blurb'].lower() == 'nan':
                spell_details.pop('short_blurb')

            os.makedirs(output_dir+f'/level_{spell_details["level"]}', exist_ok=True)

            spell_name = spell_details['name']
            if '/' in spell_name:
                spell_name = spell_name.replace('/', '-')

            # newCard = Card(spell_details, output_dir)
            # if newCard.get_has_tables():
            #     spells_with_tables.add(f"(Lvl {newCard.get_level()}) {newCard.get_name()}")
            

            log.info(f"[{count_created}/{total_count}]: Level {spell_details['level']} spell, '{spell_details['name']}' - generating...")
            create_spell_card(spell_details, output_loc=output_dir+f'/level_{spell_details["level"]}/{spell_name}.docx')
            log.debug('done.')
            
            if spell_details['has_tables']:
                spells_with_tables.add(f"(Lvl {spell_details['level']}) {spell_details['name']}")
            count_created += 1

    log.info(f"All specified cards have been generated and written to subdirectories in {output_dir}")
    if len(spells_with_tables):
        log.warning("At least one spell contained a table, you may want to check that it looks right:")
        for t in spells_with_tables:
            log.warning(f'\t* {t}')


if __name__ == "__main__":
    # Testing
    spells_df = parse_input_xlsx("spell_list_inputs.xlsx")
    # filtered_df = spells_df.iloc[19]

    # newCard = Card(filtered_df, "outputs/test2")

    # log.info(newCard.get_output_location(docx=True))
    # log.info(newCard.get_color())

    # newCard.save_as_docx()

    filtered_df = spells_df.iloc[19:65]
    print(len(filtered_df))

    create_filtered_cards(filtered_df, output_dir="outputs/test")

    
